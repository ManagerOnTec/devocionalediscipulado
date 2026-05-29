"""
Importador de conteúdo de estudos a partir de arquivos .docx.

Detecta automaticamente a estrutura do documento com base em parágrafos bold:
  DIA X — Título  (bold) → novo Tema
  Tópico          (bold) → label; próxima linha não-bold = valor do tópico
  Texto Base      (bold) → próximas linhas = texto_base do Tema
  Devocional      (bold) → próximas linhas = conteúdo de estudo
  Oração          (bold) → próximas linhas = oração
  Referências     (bold) → próximas linhas = referências cruzadas
  Conclusão       (bold) → próximas linhas = conclusão
  Exemplo         (bold) → próximas linhas = exemplo prático

Estratégias de agrupamento em Módulos:
  UNICO  → todos os temas em um único módulo
  TOPICO → cada tópico distinto vira um módulo
  GRUPOS → temas divididos em grupos de N (por ordem)
"""

import re
from io import BytesIO
from typing import Any

from docx import Document
from docx.opc.exceptions import PackageNotFoundError


# ─── Mapeamento de labels bold → campos do modelo ────────────────────────────

_SECOES = {
    "Texto Base":   "texto_base",
    "Devocional":   "estudo",
    "Oração":       "oracao",
    "Oracao":       "oracao",       # fallback sem acento
    "Referências":  "referencias_cruzadas",
    "Referencias":  "referencias_cruzadas",
    "Conclusão":    "conclusao",
    "Conclusao":    "conclusao",
    "Exemplo":      "exemplo_pratico",
}

_DIA_RE = re.compile(r"^DIA\s+(\d+)\s*[—–\-]\s*(.+)", re.IGNORECASE)


def _is_bold(paragrafo) -> bool:
    """Retorna True se algum run do parágrafo for bold."""
    return any(r.bold for r in paragrafo.runs) if paragrafo.runs else False


def _extrair_titulo_docx(doc) -> str:
    """Tenta extrair um título do documento (propriedades ou primeiro parágrafo bold)."""
    # Propriedades do documento (Core Properties)
    try:
        titulo = doc.core_properties.title
        if titulo and titulo.strip():
            return titulo.strip()
    except Exception:
        pass
    # Primeiro parágrafo não-vazio bold que NÃO seja DIA X
    for p in doc.paragraphs:
        t = p.text.strip()
        if t and _is_bold(p) and not _DIA_RE.match(t):
            return t
    return "Estudo Importado"


def _parsear_paragrafos(paragrafos: list) -> list[dict[str, Any]]:
    """
    Varre os parágrafos e retorna lista de dicts, um por DIA detectado.
    """
    temas: list[dict[str, Any]] = []
    i = 0
    n = len(paragrafos)

    while i < n:
        p = paragrafos[i]
        texto = p.text.strip()
        m = _DIA_RE.match(texto)

        if m and _is_bold(p):
            tema: dict[str, Any] = {
                "num":                int(m.group(1)),
                "titulo":             m.group(2).strip(),
                "topico":             "",
                "texto_base":         "",
                "estudo":             "",
                "oracao":             "",
                "referencias_cruzadas": "",
                "conclusao":          "",
                "exemplo_pratico":    "",
            }
            j = i + 1
            secao_atual = None
            linhas_secao: list[str] = []
            aguardando_topico = False

            while j < n:
                pj = paragrafos[j]
                tj = pj.text.strip()

                # Novo DIA → parar
                if _DIA_RE.match(tj) and _is_bold(pj):
                    break

                if _is_bold(pj) and tj:
                    # Salvar seção anterior
                    if secao_atual and linhas_secao:
                        tema[secao_atual] = "\n".join(linhas_secao).strip()
                    linhas_secao = []

                    if tj == "Tópico":
                        aguardando_topico = True
                        secao_atual = None
                    elif tj in _SECOES:
                        aguardando_topico = False
                        secao_atual = _SECOES[tj]
                    else:
                        aguardando_topico = False
                        secao_atual = None

                elif tj:  # não-bold, não-vazio
                    if aguardando_topico:
                        tema["topico"] = tj
                        aguardando_topico = False
                    elif secao_atual:
                        linhas_secao.append(tj)

                j += 1

            # Salvar última seção
            if secao_atual and linhas_secao:
                tema[secao_atual] = "\n".join(linhas_secao).strip()

            temas.append(tema)
            i = j
        else:
            i += 1

    return temas


def parsear_docx(arquivo) -> dict[str, Any]:
    """
    Ponto de entrada público.

    Parâmetros
    ----------
    arquivo : file-like object ou caminho de string

    Retorna
    -------
    dict com:
        titulo_sugerido : str
        temas           : list[dict]   — um dict por DIA detectado
        topicos_unicos  : list[str]    — tópicos distintos (preservando ordem)
        erros           : list[str]    — avisos/erros não-fatais
    """
    erros: list[str] = []

    try:
        if isinstance(arquivo, (str, bytes)):
            doc = Document(arquivo)
        else:
            conteudo = arquivo.read()
            doc = Document(BytesIO(conteudo))
    except PackageNotFoundError as exc:
        return {"titulo_sugerido": "", "temas": [], "topicos_unicos": [], "erros": [str(exc)]}
    except Exception as exc:
        return {"titulo_sugerido": "", "temas": [], "topicos_unicos": [], "erros": [f"Erro ao abrir arquivo: {exc}"]}

    titulo_sugerido = _extrair_titulo_docx(doc)
    paragrafos = doc.paragraphs

    temas = _parsear_paragrafos(list(paragrafos))

    if not temas:
        erros.append(
            "Nenhum padrão 'DIA X — Título' encontrado. "
            "Verifique se os títulos dos dias estão em negrito."
        )

    # Tópicos únicos (preservando ordem de aparição)
    topicos_unicos = list(dict.fromkeys(t["topico"] for t in temas if t["topico"]))

    return {
        "titulo_sugerido": titulo_sugerido,
        "temas": temas,
        "topicos_unicos": topicos_unicos,
        "erros": erros,
    }


def agrupar_temas(temas: list[dict], estrategia: str, tamanho_grupo: int = 5) -> list[dict]:
    """
    Agrupa a lista de temas em módulos conforme a estratégia escolhida.

    Retorna lista de dicts:
        { "titulo": str, "ordem": int, "temas": list[dict] }
    """
    if estrategia == "UNICO":
        return [{"titulo": "Módulo 1", "ordem": 1, "temas": temas}]

    if estrategia == "TOPICO":
        modulos: dict[str, list] = {}
        sem_topico: list[dict] = []
        for tema in temas:
            tp = tema["topico"] or ""
            if tp:
                modulos.setdefault(tp, []).append(tema)
            else:
                sem_topico.append(tema)
        resultado = [
            {"titulo": tp, "ordem": idx + 1, "temas": ts}
            for idx, (tp, ts) in enumerate(modulos.items())
        ]
        if sem_topico:
            resultado.append({"titulo": "Sem Tópico", "ordem": len(resultado) + 1, "temas": sem_topico})
        return resultado

    if estrategia == "GRUPOS":
        tamanho_grupo = max(1, tamanho_grupo)
        resultado = []
        for idx, inicio in enumerate(range(0, len(temas), tamanho_grupo)):
            bloco = temas[inicio : inicio + tamanho_grupo]
            resultado.append({
                "titulo": f"Módulo {idx + 1}",
                "ordem": idx + 1,
                "temas": bloco,
            })
        return resultado

    # fallback: único
    return [{"titulo": "Módulo 1", "ordem": 1, "temas": temas}]
