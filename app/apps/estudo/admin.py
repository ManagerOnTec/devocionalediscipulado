"""
Admin do app Estudo.

Registra Trilha, Modulo, Tema e ProgressoTema
com interface completa usando django-unfold + BaseModelAdmin.
"""

import json

from django import forms
from django.contrib import admin, messages
from django.http import HttpResponseRedirect
from django.shortcuts import redirect, render
from django.urls import path, reverse
from django.utils.html import format_html
from django.utils.text import slugify

from apps.core.admin import AUDIT_FIELDSET, AUDIT_READONLY_FIELDS, BaseModelAdmin, StaffAccessMixin
from unfold.admin import ModelAdmin as UnfoldModelAdmin

from django.contrib.auth import get_user_model
from django.contrib.auth.models import Group

from .importers import agrupar_temas, parsear_docx
from .models import Modulo, ProgressoTema, Tema, Trilha


# ─────────────────────────────────────────────────────────────────────────────
# Inlines
# ─────────────────────────────────────────────────────────────────────────────

class ModuloInline(StaffAccessMixin, admin.StackedInline):
    model = Modulo
    extra = 0
    fields = ("titulo", "descricao", "imagem_capa", "ordem", "acesso", "is_active")
    ordering = ("ordem",)
    show_change_link = True

    def has_delete_permission(self, request, obj=None):
        return request.user.is_superuser


class TemaInline(StaffAccessMixin, admin.StackedInline):
    model = Tema
    extra = 0
    fields = (
        "titulo", "ordem", "duracao_estimada",
        "texto_base",
        "tem_oracao", "tem_referencias", "tem_estudo", "tem_exemplo", "tem_conclusao",
        "acesso", "is_active",
    )
    ordering = ("ordem",)
    show_change_link = True

    def has_delete_permission(self, request, obj=None):
        return request.user.is_superuser


# ─────────────────────────────────────────────────────────────────────────────
# Ações em massa — Acesso
# ─────────────────────────────────────────────────────────────────────────────

@admin.action(description="🔓 Definir acesso como Público")
def acao_acesso_publico(modeladmin, request, queryset):
    total = queryset.update(acesso=Trilha.AcessoChoices.PUBLICO)
    modeladmin.message_user(request, f"{total} registro(s) definido(s) como Público.", messages.SUCCESS)


@admin.action(description="🔐 Definir acesso como Login obrigatório")
def acao_acesso_login(modeladmin, request, queryset):
    total = queryset.update(acesso=Trilha.AcessoChoices.LOGIN_OBRIGATORIO)
    modeladmin.message_user(request, f"{total} registro(s) definido(s) como Login obrigatório.", messages.SUCCESS)


@admin.action(description="🔑 Definir acesso como Permissão específica")
def acao_acesso_especifico(modeladmin, request, queryset):
    """Disponível apenas para Trilha, que suporta PERMISSAO_ESPECIFICA."""
    total = queryset.update(acesso=Trilha.AcessoChoices.PERMISSAO_ESPECIFICA)
    modeladmin.message_user(request, f"{total} trilha(s) definida(s) como Permissão específica.", messages.SUCCESS)


# ─────────────────────────────────────────────────────────────────────────────
# Trilha
# ─────────────────────────────────────────────────────────────────────────────

@admin.register(Trilha)
class TrilhaAdmin(BaseModelAdmin):
    list_display = ("titulo", "tipo", "status", "acesso", "ordem", "status_badge", "criado_em")
    list_display_links = ("titulo",)
    list_filter = ("tipo", "status", "acesso", "is_active")
    search_fields = ("titulo", "descricao")
    prepopulated_fields = {"slug": ("titulo",)}
    ordering = ("ordem", "titulo")
    inlines = [ModuloInline]
    change_list_template = "admin/estudo/trilha/change_list.html"
    actions = [acao_acesso_publico, acao_acesso_login, acao_acesso_especifico]

    fieldsets = (
        (
            "Identificação",
            {
                "fields": ("titulo", "slug", "descricao"),
            },
        ),
        (
            "Configuração",
            {
                "fields": ("tipo", "status", "acesso", "ordem"),
            },
        ),
        (
            "Controle de Acesso por Grupo",
            {
                "fields": ("grupos",),
                "description": (
                    "⚠️ Ativo somente quando Acesso = \"Permissão específica\". "
                    "Adicione os grupos que poderão acessar esta trilha. "
                    "Lembre-se de atribuir o grupo ao usuário no cadastro do usuário."
                ),
            },
        ),
    )

    class Media:
        js = ("admin/js/trilha_acesso.js",)

    def get_prepopulated_fields(self, request, obj=None):
        if obj:
            return {}
        return {"slug": ("titulo",)}

    def get_readonly_fields(self, request, obj=None):
        readonly = set(super().get_readonly_fields(request, obj))
        if obj:
            readonly.add("slug")
        return tuple(readonly)

    # ── Botão de atalho no topo da changelist ─────────────────────────────────

    def changelist_view(self, request, extra_context=None):
        extra_context = extra_context or {}
        extra_context["importar_docx_url"] = reverse("admin:estudo_trilha_importar_docx")
        return super().changelist_view(request, extra_context=extra_context)

    # ── URLs personalizadas ───────────────────────────────────────────────────

    def get_urls(self):
        urls = super().get_urls()
        custom = [
            path(
                "importar-docx/",
                self.admin_site.admin_view(self._view_importar_docx),
                name="estudo_trilha_importar_docx",
            ),
            path(
                "confirmar-import/",
                self.admin_site.admin_view(self._view_confirmar_import),
                name="estudo_trilha_confirmar_import",
            ),
        ]
        return custom + urls

    # ── View passo 1: upload + configuração ──────────────────────────────────

    def _view_importar_docx(self, request):
        ctx = {
            **self.admin_site.each_context(request),
            "title": "Importar estudo de arquivo Word",
            "opts": self.model._meta,
            "has_permission": True,
        }

        if request.method == "GET":
            return render(request, "admin/estudo/trilha/importar_docx.html", ctx)

        # POST: processar arquivo
        arquivo = request.FILES.get("arquivo")
        if not arquivo:
            messages.error(request, "Selecione um arquivo .docx.")
            return render(request, "admin/estudo/trilha/importar_docx.html", ctx)

        if not arquivo.name.lower().endswith(".docx"):
            messages.error(request, "Apenas arquivos .docx são suportados.")
            return render(request, "admin/estudo/trilha/importar_docx.html", ctx)

        resultado = parsear_docx(arquivo)

        if resultado["erros"]:
            for erro in resultado["erros"]:
                messages.warning(request, erro)

        if not resultado["temas"]:
            messages.error(request, "Nenhum tema detectado no arquivo.")
            return render(request, "admin/estudo/trilha/importar_docx.html", ctx)

        ctx.update({
            "titulo_sugerido": resultado["titulo_sugerido"],
            "total_temas": len(resultado["temas"]),
            "topicos_unicos": resultado["topicos_unicos"],
            "temas_json": json.dumps(resultado["temas"], ensure_ascii=False),
            "tipo_choices": Trilha.TipoChoices.choices,
            "status_choices": Trilha.StatusChoices.choices,
            "acesso_choices": Trilha.AcessoChoices.choices,
            "preview_temas": resultado["temas"][:5],  # mostrar primeiros 5 no preview
        })
        return render(request, "admin/estudo/trilha/confirmar_import.html", ctx)

    # ── View passo 2: criar objetos ───────────────────────────────────────────

    def _view_confirmar_import(self, request):
        if request.method != "POST":
            return redirect(reverse("admin:estudo_trilha_importar_docx"))

        # Recuperar dados do formulário
        titulo_trilha = request.POST.get("titulo_trilha", "").strip()
        tipo          = request.POST.get("tipo", Trilha.TipoChoices.DEVOCIONAL)
        status        = request.POST.get("status", Trilha.StatusChoices.RASCUNHO)
        acesso        = request.POST.get("acesso", Trilha.AcessoChoices.PUBLICO)
        agrupamento   = request.POST.get("agrupamento", "UNICO")
        tamanho_grupo = int(request.POST.get("tamanho_grupo", "5") or "5")
        temas_json    = request.POST.get("temas_json", "[]")

        if not titulo_trilha:
            messages.error(request, "O título da trilha é obrigatório.")
            return redirect(reverse("admin:estudo_trilha_importar_docx"))

        try:
            temas_raw = json.loads(temas_json)
        except json.JSONDecodeError:
            messages.error(request, "Dados inválidos. Repita o upload.")
            return redirect(reverse("admin:estudo_trilha_importar_docx"))

        # Gerar slug único para a Trilha
        slug_base = slugify(titulo_trilha)
        slug = slug_base
        contador = 2
        while Trilha.all_objects.filter(slug=slug).exists():
            slug = f"{slug_base}-{contador}"
            contador += 1

        # Criar Trilha
        trilha = Trilha.objects.create(
            titulo=titulo_trilha,
            slug=slug,
            tipo=tipo,
            status=status,
            acesso=acesso,
            criado_por=request.user,
        )

        # Agrupar temas em módulos
        modulos_dados = agrupar_temas(temas_raw, agrupamento, tamanho_grupo)

        total_modulos = 0
        total_temas = 0

        for mod_data in modulos_dados:
            # Slug único para o Módulo
            slug_mod_base = slugify(mod_data["titulo"])
            slug_mod = slug_mod_base
            c = 2
            while Modulo.all_objects.filter(slug=slug_mod).exists():
                slug_mod = f"{slug_mod_base}-{c}"
                c += 1

            modulo = Modulo.objects.create(
                trilha=trilha,
                titulo=mod_data["titulo"],
                slug=slug_mod,
                ordem=mod_data["ordem"],
                criado_por=request.user,
            )
            total_modulos += 1

            for ordem_tema, tema_data in enumerate(mod_data["temas"], start=1):
                slug_tema_base = slugify(tema_data["titulo"])
                slug_tema = slug_tema_base
                ct = 2
                while Tema.all_objects.filter(slug=slug_tema).exists():
                    slug_tema = f"{slug_tema_base}-{ct}"
                    ct += 1

                Tema.objects.create(
                    modulo=modulo,
                    titulo=tema_data["titulo"],
                    slug=slug_tema,
                    ordem=ordem_tema,
                    texto_base=tema_data.get("texto_base", ""),
                    tem_estudo=bool(tema_data.get("estudo")),
                    estudo=tema_data.get("estudo", ""),
                    tem_oracao=bool(tema_data.get("oracao")),
                    oracao=tema_data.get("oracao", ""),
                    tem_referencias=bool(tema_data.get("referencias_cruzadas")),
                    referencias_cruzadas=tema_data.get("referencias_cruzadas", ""),
                    tem_conclusao=bool(tema_data.get("conclusao")),
                    conclusao=tema_data.get("conclusao", ""),
                    tem_exemplo=bool(tema_data.get("exemplo_pratico")),
                    exemplo_pratico=tema_data.get("exemplo_pratico", ""),
                    criado_por=request.user,
                )
                total_temas += 1

        messages.success(
            request,
            f'Importação concluída! Trilha "{trilha.titulo}" criada com '
            f"{total_modulos} módulo(s) e {total_temas} tema(s).",
        )
        return HttpResponseRedirect(
            reverse("admin:estudo_trilha_change", args=[trilha.pk])
        )


# ─────────────────────────────────────────────────────────────────────────────
# Modulo
# ─────────────────────────────────────────────────────────────────────────────

@admin.register(Modulo)
class ModuloAdmin(BaseModelAdmin):
    list_display = ("titulo", "trilha", "acesso", "ordem", "status_badge", "criado_em")
    list_display_links = ("titulo",)
    list_filter = ("trilha", "acesso", "is_active")
    search_fields = ("titulo", "descricao", "trilha__titulo")
    ordering = ("trilha", "ordem")
    autocomplete_fields = ("trilha",)
    inlines = [TemaInline]
    actions = [acao_acesso_publico, acao_acesso_login]

    fieldsets = (
        (
            "Identificação",
            {
                "fields": ("trilha", "titulo", "slug", "descricao"),
            },
        ),
        (
            "Mídia e Ordenação",
            {
                "fields": ("imagem_capa", "ordem", "acesso"),
            },
        ),
    )

    def get_prepopulated_fields(self, request, obj=None):
        if obj:
            return {}
        return {"slug": ("titulo",)}

    def get_readonly_fields(self, request, obj=None):
        readonly = set(super().get_readonly_fields(request, obj))
        if obj:
            readonly.add("slug")
        return tuple(readonly)

    @admin.display(description="Imagem")
    def imagem_preview(self, obj):
        url = obj.imagem_capa_url
        if url:
            return format_html(
                '<img src="{}" style="height:40px;border-radius:4px;" alt="capa">',
                url,
            )
        return "—"


# ─────────────────────────────────────────────────────────────────────────────
# Tema
# ─────────────────────────────────────────────────────────────────────────────

@admin.register(Tema)
class TemaAdmin(BaseModelAdmin):
    list_display = (
        "titulo", "modulo", "acesso", "ordem", "duracao_estimada",
        "tem_oracao", "tem_referencias", "tem_estudo", "tem_exemplo", "tem_conclusao",
        "status_badge", "criado_em",
    )
    list_display_links = ("titulo",)
    list_filter = (
        "modulo__trilha", "modulo", "acesso",
        "tem_oracao", "tem_referencias", "tem_estudo", "tem_exemplo", "tem_conclusao",
        "is_active",
    )
    search_fields = ("titulo", "modulo__titulo", "modulo__trilha__titulo")
    ordering = ("modulo", "ordem")
    autocomplete_fields = ("modulo",)
    actions = [acao_acesso_publico, acao_acesso_login]

    fieldsets = (
        (
            "Identificação",
            {
                "fields": ("modulo", "titulo", "slug", "ordem", "duracao_estimada", "acesso"),
            },
        ),
        (
            "Conteúdo Principal",
            {
                "fields": ("texto_base",),
            },
        ),
        (
            "Oração",
            {
                "fields": ("tem_oracao", "oracao"),
                "classes": ("collapse",),
                "description": "Ative para exibir uma oração ao final do tema.",
            },
        ),
        (
            "Referências Cruzadas",
            {
                "fields": ("tem_referencias", "referencias_cruzadas"),
                "classes": ("collapse",),
            },
        ),
        (
            "Estudo Aprofundado",
            {
                "fields": ("tem_estudo", "estudo"),
                "classes": ("collapse",),
            },
        ),
        (
            "Exemplo Prático",
            {
                "fields": ("tem_exemplo", "exemplo_pratico"),
                "classes": ("collapse",),
            },
        ),
        (
            "Conclusão",
            {
                "fields": ("tem_conclusao", "conclusao"),
                "classes": ("collapse",),
            },
        ),
    )

    def get_prepopulated_fields(self, request, obj=None):
        if obj:
            return {}
        return {"slug": ("titulo",)}

    def get_readonly_fields(self, request, obj=None):
        readonly = set(super().get_readonly_fields(request, obj))
        if obj:
            readonly.add("slug")
        return tuple(readonly)


# ─────────────────────────────────────────────────────────────────────────────
# ProgressoTema
# ─────────────────────────────────────────────────────────────────────────────

@admin.register(ProgressoTema)
class ProgressoTemaAdmin(StaffAccessMixin, UnfoldModelAdmin):
    """Uma linha por usuário × módulo — exibe o último tema concluído e data."""

    change_list_template = "admin/estudo/progressotema/change_list.html"

    list_display = (
        "usuario",
        "trilha_col",
        "modulo_col",
        "percentual_modulo",
        "ultimo_tema_col",
        "data_conclusao",
    )
    list_filter = ("tema__modulo__trilha", "tema__modulo")
    search_fields = ("usuario__email", "tema__titulo", "tema__modulo__titulo", "tema__modulo__trilha__titulo")
    ordering = ("usuario__nome_completo", "tema__modulo__trilha__titulo", "tema__modulo__titulo")
    readonly_fields = ("criado_em", "atualizado_em", "data_conclusao")

    def _qs_usuarios_visiveis(self, request):
        """
        Retorna PKs de usuários cujos progressos o admin logado pode ver.
        None = sem restrição (superusuário ou staff com acesso admin).
        """
        if request.user.is_staff:
            return None
        # Usuário comum: apenas o próprio progresso
        return [request.user.pk]

    def get_queryset(self, request):
        from django.db.models import Max

        # Uma linha por (usuário, módulo): o ProgressoTema com o maior pk = mais recente
        latest_ids = (
            ProgressoTema.objects
            .values("usuario_id", "tema__modulo_id")
            .annotate(_latest=Max("id"))
            .values("_latest")
        )
        qs = (
            super()
            .get_queryset(request)
            .filter(id__in=latest_ids)
            .select_related("usuario", "tema__modulo__trilha")
        )
        visiveis = self._qs_usuarios_visiveis(request)
        if visiveis is not None:
            qs = qs.filter(usuario_id__in=visiveis)
        return qs

    # ── Colunas calculadas ────────────────────────────────────────────────────

    @admin.display(description="Trilha", ordering="tema__modulo__trilha__titulo")
    def trilha_col(self, obj):
        trilha = obj.tema.modulo.trilha
        url = reverse("admin:estudo_trilha_change", args=[trilha.pk])
        return format_html('<a href="{}">{}</a>', url, trilha.titulo)

    @admin.display(description="Módulo", ordering="tema__modulo__titulo")
    def modulo_col(self, obj):
        modulo = obj.tema.modulo
        url = reverse("admin:estudo_modulo_change", args=[modulo.pk])
        return format_html('<a href="{}">{}</a>', url, modulo.titulo)

    @admin.display(description="Último tema", ordering="tema__titulo")
    def ultimo_tema_col(self, obj):
        url = reverse("admin:estudo_tema_change", args=[obj.tema.pk])
        return format_html('<a href="{}">{}</a>', url, obj.tema.titulo)

    @admin.display(description="% Módulo", ordering="tema__modulo")
    def percentual_modulo(self, obj):
        """Percentual de temas concluídos pelo usuário no módulo deste tema."""
        modulo = obj.tema.modulo
        total = modulo.temas.count()
        if not total:
            return "—"
        concluidos = ProgressoTema.objects.filter(
            usuario=obj.usuario,
            tema__modulo=modulo,
        ).count()
        pct = int(concluidos / total * 100)
        cor = "#198754" if pct == 100 else "#0d6efd"
        return format_html(
            '<div style="display:flex;align-items:center;gap:6px;min-width:120px;">'
            '<div style="flex:1;height:8px;background:#e9ecef;border-radius:4px;overflow:hidden;">'
            '<div style="width:{}%;height:100%;background:{};border-radius:4px;"></div>'
            "</div>"
            '<span style="font-size:.75rem;font-weight:600;color:{};">{}%&nbsp;({}/{})</span>'
            "</div>",
            pct, cor, cor, pct, concluidos, total,
        )

    # ── Changelist com indicadores KPI ────────────────────────────────────────

    def changelist_view(self, request, extra_context=None):
        from django.db.models import Count
        from django.utils import timezone

        from .models import Tema as _Tema

        extra_context = extra_context or {}
        extra_context["title"] = "Progressos por Módulo"

        # Escopa os KPIs pela mesma regra de visibilidade do get_queryset
        visiveis = self._qs_usuarios_visiveis(request)
        qs_base = ProgressoTema.objects.all()
        if visiveis is not None:
            qs_base = qs_base.filter(usuario_id__in=visiveis)

        total_alunos = qs_base.values("usuario").distinct().count()
        total_conclusoes = qs_base.count()
        hoje = timezone.localdate()
        conclusoes_hoje = qs_base.filter(data_conclusao__date=hoje).count()

        # Percentual médio: para cada par (usuário, módulo) calcula % e tira a média
        pares = list(
            qs_base
            .values("usuario", "tema__modulo")
            .annotate(concluidos=Count("id"))
        )
        modulo_totais = dict(
            _Tema.objects
            .values("modulo_id")
            .annotate(total=Count("id"))
            .values_list("modulo_id", "total")
        )
        soma_pct = sum(
            par["concluidos"] / modulo_totais[par["tema__modulo"]] * 100
            for par in pares
            if modulo_totais.get(par["tema__modulo"])
        )
        n_pares = sum(1 for par in pares if modulo_totais.get(par["tema__modulo"]))
        pct_medio = round(soma_pct / n_pares, 1) if n_pares else 0

        extra_context["kpi_progresso"] = [
            {
                "title": "Alunos com progresso",
                "metric": str(total_alunos),
                "icon": "school",
                "description": "Usuários com ao menos um tema concluído",
            },
            {
                "title": "Temas concluídos",
                "metric": str(total_conclusoes),
                "icon": "task_alt",
                "description": "Total de conclusões registradas",
            },
            {
                "title": "Percentual médio",
                "metric": f"{pct_medio}%",
                "icon": "trending_up",
                "description": "Média de conclusão por módulo/aluno",
            },
            {
                "title": "Conclusões hoje",
                "metric": str(conclusoes_hoje),
                "icon": "today",
                "description": f"Registros em {hoje.strftime('%d/%m/%Y')}",
            },
        ]
        return super().changelist_view(request, extra_context=extra_context)

    def get_fieldsets(self, request, obj=None):
        return [
            (
                None,
                {
                    "fields": ("usuario", "tema", "data_conclusao"),
                },
            ),
            (
                "Auditoria",
                {
                    "fields": ("criado_em", "atualizado_em"),
                    "classes": ("collapse",),
                },
            ),
        ]

    def formfield_for_foreignkey(self, db_field, request, **kwargs):
        """Restringe o campo usuario a usuários ativos."""
        User = get_user_model()
        if db_field.related_model is User:
            kwargs["queryset"] = User.objects.filter(is_active=True).order_by("nome_completo")
        return super().formfield_for_foreignkey(db_field, request, **kwargs)

    def get_readonly_fields(self, request, obj=None):
        return self.readonly_fields
    def has_delete_permission(self, request, obj=None):
        return request.user.is_superuser


# ─────────────────────────────────────────────────────────────────────────────
# Group — admin customizado (sem permissões Django; mostra trilhas vinculadas)
# ─────────────────────────────────────────────────────────────────────────────


class TrilhaGrupoInline(StaffAccessMixin, admin.TabularInline):
    """Mostra e gerencia as trilhas de acesso restrito vinculadas a este grupo."""

    model = Trilha.grupos.through
    extra = 1
    verbose_name = "Trilha liberada"
    verbose_name_plural = "Trilhas liberadas para este grupo"
    fk_name = "group"
    autocomplete_fields = ("trilha",)

    def get_queryset(self, request):
        return (
            super()
            .get_queryset(request)
            .select_related("trilha")
        )

    def has_delete_permission(self, request, obj=None):
        return request.user.is_superuser


admin.site.unregister(Group)


@admin.register(Group)
class GrupoAdmin(StaffAccessMixin, UnfoldModelAdmin):
    """
    Admin de Grupo sem exposição de permissões Django.
    Exibe inline com as Trilhas de acesso restrito que este grupo libera.
    """

    search_fields = ("name",)
    ordering = ("name",)
    inlines = [TrilhaGrupoInline]

    compressed_fields = True
    warn_unsaved_changes = True

    def get_fieldsets(self, request, obj=None):
        return [
            (
                None,
                {
                    "fields": ("name",),
                },
            ),
        ]

    def has_delete_permission(self, request, obj=None):
        return request.user.is_superuser


# ─────────────────────────────────────────────────────────────────────────────
# EstudoPessoal — acesso exclusivo para superadmin
# ─────────────────────────────────────────────────────────────────────────────

from .models import EstudoPessoal, TopicoEstudo


# ── Helpers de exportação ────────────────────────────────────────────────────

def _linhas_estudo(obj) -> list:
    """
    Retorna lista de (titulo_secao, conteudo) para exportação,
    respeitando os booleans de inclusão.
    """
    secoes = []

    def add(titulo, conteudo):
        """Só adiciona a seção se o conteúdo for não-vazio."""
        conteudo = (conteudo or "").strip()
        if conteudo:
            secoes.append((titulo, conteudo))

    add("REFERÊNCIA BÍBLICA", obj.referencia)
    add("TEXTO BÍBLICO", obj.texto_biblico)
    add("CONTEXTO ANTERIOR", obj.contexto_anterior)
    add("CONTEXTO POSTERIOR", obj.contexto_posterior)
    add("VERSÍCULOS RELACIONADOS", obj.versiculos_relacionados)
    add("PALAVRA CENTRAL", obj.palavra_central)
    add("PALAVRA QUE SE REPETE", obj.palavra_repetida)
    add("PALAVRA A APROFUNDAR", obj.palavra_aprofundar)
    add("SINÔNIMOS", obj.sinonimos)
    add("TRADUÇÃO / HEBRAICO", obj.traducao_hebraico)
    add("TRADUÇÃO / GREGO", obj.traducao_grego)
    add("TRADUÇÃO / LATIM", obj.traducao_latim)
    add("TRADUÇÃO / INGLÊS", obj.traducao_ingles)
    add("OBSERVAÇÕES PARA O PORTUGUÊS", obj.observacoes_portugues)
    add("ONDE ESTÁ CRISTO", obj.onde_esta_cristo)
    add("USO NO ANTIGO TESTAMENTO", obj.uso_antigo_testamento)
    add("USO NO NOVO TESTAMENTO", obj.uso_novo_testamento)
    add("USO HISTÓRICO GREGO", obj.uso_historico_grego)
    add("USO NA FILOSOFIA ANTIGA", obj.uso_historico_filosofia_antiga)
    add("FILOSOFIA MODERNA", obj.filosofia_moderna)
    add("GÊNERO LITERÁRIO", obj.genero_literario)
    add("CULTURA", obj.cultura)
    add("INTENÇÃO DO AUTOR", obj.intencao_autor)
    add("QUEM ESTÁ FALANDO?", obj.quem_fala)
    add("PARA QUEM?", obj.para_quem)
    add("SOBRE O QUÊ?", obj.sobre_o_que)
    add("QUAL O OBJETIVO?", obj.qual_objetivo)
    add("O QUE EXIGE DE MIM?", obj.o_que_exige_de_mim)

    if obj.incluir_introducao:
        add("INTRODUÇÃO", obj.introducao)

    topicos = obj.topicos.filter(incluir=True).order_by("ordem")
    if topicos.exists():
        secoes.append(("APLICAÇÕES / TÓPICOS", ""))
        for t in topicos:
            secoes.append((f"  Tópico {t.ordem}: {t.titulo}", (t.conteudo or "").strip()))

    if obj.incluir_explicacao:
        add("EXPLICAÇÃO", obj.explicacao)
    if obj.incluir_aplicacao:
        add("APLICAÇÃO", obj.aplicacao)
    if obj.incluir_conclusao:
        add("CONCLUSÃO", obj.conclusao)
    if obj.incluir_oracao:
        add("ORAÇÃO", obj.oracao)

    return secoes


def _exportar_txt(obj):
    """Gera HttpResponse com arquivo TXT."""
    from django.http import HttpResponse

    linhas = [f"{'=' * 60}", f"  {obj.titulo.upper()}", f"  {obj.referencia}", f"{'=' * 60}", ""]
    for titulo, conteudo in _linhas_estudo(obj):
        linhas.append(f"\n{'─' * 50}")
        linhas.append(f"  {titulo}")
        linhas.append(f"{'─' * 50}")
        if conteudo:
            linhas.append(conteudo)
        linhas.append("")

    response = HttpResponse("\n".join(linhas), content_type="text/plain; charset=utf-8")
    nome = f"estudo_{obj.pk}_{obj.referencia[:30].replace(' ', '_')}.txt"
    response["Content-Disposition"] = f'attachment; filename="{nome}"'
    return response


def _exportar_pdf(obj):
    """Gera HttpResponse com arquivo PDF via reportlab."""
    import io
    from django.http import HttpResponse
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=A4,
        leftMargin=2.5 * cm, rightMargin=2.5 * cm,
        topMargin=2.5 * cm, bottomMargin=2.5 * cm,
    )
    styles = getSampleStyleSheet()
    s_titulo = ParagraphStyle("Titulo", parent=styles["Heading1"], fontSize=16, spaceAfter=6, leading=22)
    s_ref = ParagraphStyle("Ref", parent=styles["Normal"], fontSize=11, spaceAfter=12, textColor=(0.3, 0.3, 0.3))
    s_h2 = ParagraphStyle("H2", parent=styles["Heading2"], fontSize=11, spaceBefore=14, spaceAfter=4, leading=15)
    s_body = ParagraphStyle("Body", parent=styles["Normal"], fontSize=10, leading=14, spaceAfter=6)

    story = [Paragraph(obj.titulo, s_titulo), Paragraph(obj.referencia, s_ref)]
    for titulo, conteudo in _linhas_estudo(obj):
        story.append(Paragraph(titulo, s_h2))
        if conteudo:
            for par in conteudo.split("\n"):
                par = par.strip()
                if par:
                    story.append(Paragraph(par, s_body))
        story.append(Spacer(1, 4))

    doc.build(story)
    pdf_bytes = buffer.getvalue()
    buffer.close()

    response = HttpResponse(pdf_bytes, content_type="application/pdf")
    nome = f"estudo_{obj.pk}_{obj.referencia[:30].replace(' ', '_')}.pdf"
    response["Content-Disposition"] = f'attachment; filename="{nome}"'
    return response


def _exportar_docx(obj):
    """Gera HttpResponse com arquivo DOCX via python-docx."""
    import io
    from django.http import HttpResponse
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document()
    h1 = doc.add_heading(obj.titulo, level=1)
    h1.runs[0].font.size = Pt(18)
    ref_par = doc.add_paragraph(obj.referencia)
    ref_par.runs[0].font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    ref_par.runs[0].font.size = Pt(12)
    doc.add_paragraph("")

    for titulo, conteudo in _linhas_estudo(obj):
        doc.add_heading(titulo, level=2)
        if conteudo:
            for linha in conteudo.split("\n"):
                linha = linha.strip()
                if linha:
                    doc.add_paragraph(linha)

    buffer = io.BytesIO()
    doc.save(buffer)
    docx_bytes = buffer.getvalue()
    buffer.close()

    response = HttpResponse(
        docx_bytes,
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    nome = f"estudo_{obj.pk}_{obj.referencia[:30].replace(' ', '_')}.docx"
    response["Content-Disposition"] = f'attachment; filename="{nome}"'
    return response


# ── Ações em massa ────────────────────────────────────────────────────────────

@admin.action(description="📄 Exportar como TXT")
def acao_exportar_txt(modeladmin, request, queryset):
    if queryset.count() == 1:
        return _exportar_txt(queryset.first())
    modeladmin.message_user(request, "Selecione apenas 1 estudo por vez para exportar.", messages.WARNING)


@admin.action(description="📕 Exportar como PDF")
def acao_exportar_pdf(modeladmin, request, queryset):
    if queryset.count() == 1:
        return _exportar_pdf(queryset.first())
    modeladmin.message_user(request, "Selecione apenas 1 estudo por vez para exportar.", messages.WARNING)


@admin.action(description="📝 Exportar como DOCX")
def acao_exportar_docx(modeladmin, request, queryset):
    if queryset.count() == 1:
        return _exportar_docx(queryset.first())
    modeladmin.message_user(request, "Selecione apenas 1 estudo por vez para exportar.", messages.WARNING)


# ── Inline de tópicos ─────────────────────────────────────────────────────────

class TopicoEstudoInline(admin.TabularInline):
    model = TopicoEstudo
    extra = 1
    fields = ("ordem", "incluir", "titulo", "conteudo")
    ordering = ("ordem",)


# ── Admin EstudoPessoal ───────────────────────────────────────────────────────

@admin.register(EstudoPessoal)
class EstudoPessoalAdmin(UnfoldModelAdmin):
    """Visível e editável somente por superadmin."""

    list_display = ("titulo", "referencia", "criado_em", "atualizado_em")
    list_display_links = ("titulo",)
    search_fields = ("titulo", "referencia")
    ordering = ("-criado_em",)
    readonly_fields = ("criado_em", "atualizado_em")
    inlines = [TopicoEstudoInline]
    actions = [acao_exportar_txt, acao_exportar_pdf, acao_exportar_docx]

    compressed_fields = True
    warn_unsaved_changes = True

    fieldsets = [
        ("Identificação", {
            "fields": ("titulo", "referencia", "criado_em", "atualizado_em"),
        }),
        ("Texto e Contexto Bíblico", {
            "fields": ("texto_biblico", "contexto_anterior", "contexto_posterior", "versiculos_relacionados"),
            "classes": ("collapse",),
        }),
        ("Palavras-chave e Léxico", {
            "fields": (
                "palavra_central", "palavra_repetida", "palavra_aprofundar", "sinonimos",
                "traducao_hebraico", "traducao_grego", "traducao_latim", "traducao_ingles",
                "observacoes_portugues",
            ),
            "classes": ("collapse",),
        }),
        ("Teologia e Contexto Histórico", {
            "fields": (
                "onde_esta_cristo",
                "uso_antigo_testamento", "uso_novo_testamento",
                "uso_historico_grego", "uso_historico_filosofia_antiga",
                "filosofia_moderna", "genero_literario", "cultura", "intencao_autor",
            ),
            "classes": ("collapse",),
        }),
        ("Perguntas Hermenêuticas", {
            "fields": ("quem_fala", "para_quem", "sobre_o_que", "qual_objetivo", "o_que_exige_de_mim"),
            "classes": ("collapse",),
        }),
        ("Desenvolvimento — 1. Introdução", {
            "fields": ("incluir_introducao", "introducao"),
            "description": "Abertura do estudo: gancho, problema, relevância.",
            "classes": ("collapse",),
        }),
        ("Desenvolvimento — 2. Aplicações / Tópicos", {
            "fields": (),
            "description": (
                "⬇️  Os tópicos ficam na seção 'Tópicos' abaixo deste formulário — "
                "na exportação (TXT / PDF / DOCX) e no template eles são inseridos "
                "automaticamente nesta posição, entre a Introdução e a Explicação."
            ),
        }),
        ("Desenvolvimento — 3. Explicação", {
            "fields": ("incluir_explicacao", "explicacao"),
            "description": "Exegese e exposição versículo a versículo.",
            "classes": ("collapse",),
        }),
        ("Desenvolvimento — 4. Aplicação", {
            "fields": ("incluir_aplicacao", "aplicacao"),
            "description": "Como este texto deve transformar a vida do ouvinte?",
            "classes": ("collapse",),
        }),
        ("Desenvolvimento — 5. Conclusão", {
            "fields": ("incluir_conclusao", "conclusao"),
            "description": "Síntese, chamada à ação ou convite.",
            "classes": ("collapse",),
        }),
        ("Desenvolvimento — 6. Oração", {
            "fields": ("incluir_oracao", "oracao"),
            "description": "Sugestão de oração para encerrar.",
            "classes": ("collapse",),
        }),
    ]

    # ── Restrição total a superadmin ─────────────────────────────────────────

    def has_module_perms(self, request):
        return request.user.is_superuser

    def has_view_permission(self, request, obj=None):
        return request.user.is_superuser

    def has_add_permission(self, request):
        return request.user.is_superuser

    def has_change_permission(self, request, obj=None):
        return request.user.is_superuser

    def has_delete_permission(self, request, obj=None):
        return request.user.is_superuser

    # ── URLs customizadas para exportação por botão ───────────────────────────

    def get_urls(self):
        from django.urls import path as dj_path
        urls = super().get_urls()
        custom = [
            dj_path(
                "<int:pk>/exportar/txt/",
                self.admin_site.admin_view(self._view_exportar_txt),
                name="estudo_estudopessoal_exportar_txt",
            ),
            dj_path(
                "<int:pk>/exportar/pdf/",
                self.admin_site.admin_view(self._view_exportar_pdf),
                name="estudo_estudopessoal_exportar_pdf",
            ),
            dj_path(
                "<int:pk>/exportar/docx/",
                self.admin_site.admin_view(self._view_exportar_docx),
                name="estudo_estudopessoal_exportar_docx",
            ),
        ]
        return custom + urls

    def _get_obj_or_403(self, request, pk):
        from django.core.exceptions import PermissionDenied
        if not request.user.is_superuser:
            raise PermissionDenied
        return EstudoPessoal.objects.get(pk=pk)

    def _view_exportar_txt(self, request, pk):
        return _exportar_txt(self._get_obj_or_403(request, pk))

    def _view_exportar_pdf(self, request, pk):
        return _exportar_pdf(self._get_obj_or_403(request, pk))

    def _view_exportar_docx(self, request, pk):
        return _exportar_docx(self._get_obj_or_403(request, pk))

    def change_view(self, request, object_id, form_url="", extra_context=None):
        extra_context = extra_context or {}
        extra_context["exportar_urls"] = {
            "txt":  "../exportar/txt/",
            "pdf":  "../exportar/pdf/",
            "docx": "../exportar/docx/",
        }
        return super().change_view(request, object_id, form_url, extra_context)
